import io
import os
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

TEMPLATES_DIR = Path(__file__).parent.parent.parent / "templates"
ASSETS_DIR    = Path(__file__).parent.parent.parent / "assets"
SIGNATURE_PATH = ASSETS_DIR / "signature.png"

# ── Форматировщики ────────────────────────────────────────────────────────────

_MONTHS_RU = [
    "", "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]


def format_date_full(date_str: str) -> str:
    """'17.11.2024' → '«17» ноября 2024 г.'"""
    try:
        d, m, y = date_str.strip().split(".")
        return f"«{d}» {_MONTHS_RU[int(m)]} {y} г."
    except Exception:
        return date_str


def format_date_part(date_str: str) -> str:
    """'17.11.2024' → '17» ноября 2024 г.'  (без открывающей «)"""
    try:
        d, m, y = date_str.strip().split(".")
        return f"{d}» {_MONTHS_RU[int(m)]} {y} г."
    except Exception:
        return date_str


def format_date_legal(date_str: str) -> str:
    """'05.05.2025' → ' «05» мая 2025 г.'  (с пробелом в начале, как в шаблоне)"""
    try:
        d, m, y = date_str.strip().split(".")
        return f" «{d}» {_MONTHS_RU[int(m)]} {y} г."
    except Exception:
        return date_str


def format_amount_rub(amount: float) -> str:
    """142000.0 → '142 000 руб. 00 коп.'"""
    rubles = int(amount)
    kopecks = round((amount - rubles) * 100)
    return f"{rubles:,} руб. {kopecks:02d} коп.".replace(",", " ")


_ONES = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
_ONES_F = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
_TENS = ["", "десять", "двадцать", "тридцать", "сорок", "пятьдесят",
         "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
_HUNDREDS = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
             "шестьсот", "семьсот", "восемьсот", "девятьсот"]
_TEENS = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
          "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]


def _chunk_words(n: int, feminine: bool = False) -> str:
    ones = _ONES_F if feminine else _ONES
    parts = []
    h = n // 100
    rest = n % 100
    if h:
        parts.append(_HUNDREDS[h])
    if 10 <= rest <= 19:
        parts.append(_TEENS[rest - 10])
    else:
        t = rest // 10
        o = rest % 10
        if t:
            parts.append(_TENS[t])
        if o:
            parts.append(ones[o])
    return " ".join(parts)


def _millions(n: int) -> str:
    if n == 0:
        return ""
    last2 = n % 100
    last1 = n % 10
    word = _chunk_words(n, feminine=True)
    if 11 <= last2 <= 19:
        suffix = "миллионов"
    elif last1 == 1:
        suffix = "миллион"
    elif 2 <= last1 <= 4:
        suffix = "миллиона"
    else:
        suffix = "миллионов"
    return f"{word} {suffix}"


def _thousands(n: int) -> str:
    if n == 0:
        return ""
    last2 = n % 100
    last1 = n % 10
    word = _chunk_words(n, feminine=True)
    if 11 <= last2 <= 19:
        suffix = "тысяч"
    elif last1 == 1:
        suffix = "тысяча"
    elif 2 <= last1 <= 4:
        suffix = "тысячи"
    else:
        suffix = "тысяч"
    return f"{word} {suffix}"


def num_to_words(amount: float) -> str:
    """142000 → 'сто сорок две тысячи рублей'"""
    n = int(round(amount))
    if n == 0:
        return "ноль рублей"
    parts = []
    mils = n // 1_000_000
    rest = n % 1_000_000
    ths  = rest // 1000
    ones = rest % 1000
    if mils:
        parts.append(_millions(mils))
    if ths:
        parts.append(_thousands(ths))
    if ones:
        last2 = ones % 100
        last1 = ones % 10
        word = _chunk_words(ones)
        if 11 <= last2 <= 19:
            suffix = "рублей"
        elif last1 == 1:
            suffix = "рубль"
        elif 2 <= last1 <= 4:
            suffix = "рубля"
        else:
            suffix = "рублей"
        parts.append(f"{word} {suffix}")
    else:
        parts.append("рублей")
    return " ".join(parts)


def format_passport_series_number(series: str, number: str) -> str:
    """'1111', '111111' → '1111 № 111111'"""
    return f"{series} № {number}"


def format_passport_long(series: str, number: str) -> str:
    """'1111', '111111' → 'Паспорт: серия 1111  номер 111111'"""
    return f"Паспорт: серия {series}  номер {number}"


def format_issued_by_and_date(issued_by: str, issue_date: str) -> str:
    """Форматирует строку для согласия: 'Отделением... «15» января 2016 г.'"""
    date_full = format_date_full(issue_date)
    return f"{issued_by} {date_full}"


# ── Замена плейсхолдеров ──────────────────────────────────────────────────────

def _replace_in_paragraph(paragraph, replacements: dict) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    changed = False
    for key, value in replacements.items():
        ph = f"{{{{{key}}}}}"
        if ph in full_text:
            full_text = full_text.replace(ph, str(value) if value is not None else "")
            changed = True
    if not changed or not paragraph.runs:
        return

    parts = full_text.split('\n')
    paragraph.runs[0].text = parts[0]
    for run in paragraph.runs[1:]:
        run.text = ""

    if len(parts) > 1:
        run0 = paragraph.runs[0]
        for part in parts[1:]:
            br = run0._r.makeelement(qn('w:br'))
            run0._r.append(br)
            t_elem = run0._r.makeelement(qn('w:t'))
            t_elem.text = part
            if part != part.strip():
                t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            run0._r.append(t_elem)


def _insert_executor_signature(doc: Document) -> None:
    """Вставляет подпись исполнителя во все таблицы подписей где упоминается исполнитель."""
    if not SIGNATURE_PATH.exists():
        return

    for table in doc.tables:
        table_text = "\n".join(
            p.text
            for row in table.rows
            for cell in row.cells
            for p in cell.paragraphs
        )
        is_executor_table = any(x in table_text for x in ("Распопов", "предприниматель", "Исполнитель"))
        if not is_executor_table:
            continue

        for row in table.rows:
            if not row.cells:
                continue
            cell = row.cells[0]
            cell_text = "\n".join(p.text for p in cell.paragraphs)
            if "_____" not in cell_text:
                continue

            for para in cell.paragraphs:
                if "_____" not in para.text:
                    continue
                for run in para.runs:
                    run.text = (
                        run.text
                        .replace("________________________/_____________/", "")
                        .replace("________________________", "")
                        .replace("/_____________/", "")
                        .replace("_____________", "")
                    )
                para.add_run().add_picture(str(SIGNATURE_PATH), width=Inches(2))
                break


def _remove_empty_list_paragraphs(doc: Document) -> None:
    """Удаляет пустые параграфы с list-форматированием (иначе рендерятся как '-')."""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    numPr_tag = f"{{{ns}}}numPr"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                to_remove = [
                    p._element for p in paragraphs[:-1]
                    if not p.text.strip() and p._element.find(f".//{numPr_tag}") is not None
                ]
                for el in to_remove:
                    el.getparent().remove(el)


def _add_page_breaks(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Приложение №") or text == "СОГЛАСИЕ":
            paragraph.paragraph_format.page_break_before = True


def _fill_document(template_path: Path, replacements: dict) -> bytes:
    doc = Document(str(template_path))

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    _remove_empty_list_paragraphs(doc)
    _add_page_breaks(doc)
    _insert_executor_signature(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_to_pdf(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "contract.docx")
        pdf_path  = os.path.join(tmpdir, "contract.pdf")
        profile_dir = os.path.join(tmpdir, "loprofile")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        subprocess.run(
            [
                "libreoffice", "--headless", "--norestore",
                f"-env:UserInstallation=file://{profile_dir}",
                "--convert-to", "pdf", "--outdir", tmpdir, docx_path,
            ],
            check=True, capture_output=True, timeout=120,
        )
        with open(pdf_path, "rb") as f:
            return f.read()


# ── Сборка плейсхолдеров для физлица ─────────────────────────────────────────

def build_individual_replacements(data: dict) -> dict:
    """Собирает словарь плейсхолдеров из данных флоу физлица."""
    ru  = data["ru_passport"]
    fg  = data["foreign_passport"]
    tourists = data.get("tourists", [])
    number   = str(data["contract_number"])
    date_str = data.get("contract_date", "")

    total    = float(data.get("total_price", 0))
    deposit  = float(data.get("deposit", 0))
    remaining = float(data.get("remaining", 0))

    # Туристы (до 2 в таблице; остальные — в ADD_COND)
    def tourist_row(t):
        return {
            "name":        f"{t.get('surname_latin','')} {t.get('name_latin','')}".strip(),
            "gender":      t.get("gender", ""),
            "dob":         t.get("date_of_birth", ""),
            "passport":    t.get("passport_number", ""),
            "valid_until": t.get("valid_until", ""),
        }

    t1 = tourist_row(tourists[0]) if len(tourists) > 0 else {}
    t2 = tourist_row(tourists[1]) if len(tourists) > 1 else {}

    # Дополнительные условия — все в ADD_COND_1 через переносы строк
    add_cond_raw = data.get("additional_conditions", "") or ""
    add_cond_lines = [l.strip() for l in add_cond_raw.splitlines() if l.strip()]
    combined_cond = "\n".join(f"- {line}" for line in add_cond_lines) if add_cond_lines else ""
    add_cond = {"ADD_COND_1": combined_cond, **{f"ADD_COND_{i+1}": "" for i in range(1, 8)}}

    # Количество номеров
    room_count = data.get("room_count", "")
    room_count_display = f"{room_count} (при двухместном размещении)" if room_count == "½" else room_count

    return {
        "CONTRACT_NUMBER_FULL":          f"№ {number} ",
        "CONTRACT_NUMBER_SHORT":         number,
        "CONTRACT_DATE":                 format_date_full(date_str),
        "CONTRACT_DATE_PART":            format_date_part(date_str),
        "CLIENT_FULL_NAME":              ru.get("full_name", ""),
        "CANCELLATION_FEE_1":            "50%",
        "CANCELLATION_DAYS_1":           "30",
        "CANCELLATION_DAYS_2":           "30",
        "PASSPORT_SERIES_NUMBER":        format_passport_series_number(ru.get("series",""), ru.get("number","")),
        "PASSPORT_ISSUED_BY":            ru.get("issued_by", ""),
        "PASSPORT_ISSUE_DATE":           ru.get("issue_date", ""),
        "REGISTRATION_ADDRESS":          ru.get("registration_address", ""),
        "CLIENT_PHONE":                  data.get("phone", ""),
        "CLIENT_EMAIL":                  data.get("email", ""),
        "TOURIST_1_NAME":                t1.get("name", ""),
        "TOURIST_1_GENDER":              t1.get("gender", ""),
        "TOURIST_1_DOB":                 t1.get("dob", ""),
        "TOURIST_1_PASSPORT":            t1.get("passport", ""),
        "TOURIST_1_VALID_UNTIL":         t1.get("valid_until", ""),
        "TOURIST_2_NAME":                t2.get("name", ""),
        "TOURIST_2_GENDER":              t2.get("gender", ""),
        "TOURIST_2_DOB":                 t2.get("dob", ""),
        "TOURIST_2_PASSPORT":            t2.get("passport", ""),
        "TOURIST_2_VALID_UNTIL":         t2.get("valid_until", ""),
        "COUNTRY_CITY":                  f"{data.get('country','')} {data.get('city','')}".strip(", "),
        "HOTEL":                         data.get("hotel", ""),
        "CHECK_IN_DATE":                 data.get("check_in_date", ""),
        "CHECK_OUT_DATE":                data.get("check_out_date", ""),
        "NIGHTS":                        str(data.get("nights", "")),
        "ROOM_TYPE":                     data.get("room_type", ""),
        "ROOM_COUNT":                    room_count_display,
        "MEAL_TYPE":                     data.get("meal_type", ""),
        "TRANSFER":                      data.get("transfer", "нет"),
        "INSURANCE":                     "да" if data.get("insurance") else "нет",
        **add_cond,
        "TOTAL_PRICE_RUB":               format_amount_rub(total),
        "TOTAL_PRICE_WORDS":             num_to_words(total),
        "DEPOSIT_RUB":                   format_amount_rub(deposit),
        "DEPOSIT_WORDS":                 num_to_words(deposit),
        "DEPOSIT_DATE":                  format_date_full(date_str),
        "PAYMENT_DEADLINE":              data.get("finance_payment_deadline", ""),
        "REMAINING_RUB":                 format_amount_rub(remaining),
        "REMAINING_WORDS":               num_to_words(remaining),
        "PASSPORT_SERIES_NUMBER_LONG":   format_passport_long(ru.get("series",""), ru.get("number","")),
        "PASSPORT_ISSUED_BY_AND_DATE":   format_issued_by_and_date(ru.get("issued_by",""), ru.get("issue_date","")),
    }


# ── Сборка плейсхолдеров для юрлица ──────────────────────────────────────────

def build_legal_replacements(data: dict) -> dict:
    """Собирает словарь плейсхолдеров из данных флоу юрлица."""
    c         = data["company"]
    employees = data.get("employees", [])
    number    = str(data["contract_number"])
    date_str  = data.get("contract_date", "")
    total     = float(data.get("total_price", 0))

    def emp_row(e):
        return {
            "name":       f"{e.get('surname_latin','')} {e.get('name_latin','')}".strip(),
            "dob":        e.get("date_of_birth", ""),
            "passport":   e.get("passport_number", ""),
            "issue_date": e.get("issue_date", ""),
            "valid_until":e.get("valid_until", ""),
        }

    e1 = emp_row(employees[0]) if len(employees) > 0 else {}
    e2 = emp_row(employees[1]) if len(employees) > 1 else {}

    # Юридическая форма и короткое имя: "ООО «Торговый Дом Батиссур»" → form="ООО", short="«Торговый Дом Батиссур»"
    full_name = c.get("company_name", "")
    legal_form = c.get("legal_form", "")
    short_name_raw = full_name.replace(legal_form, "").strip()
    short_name_display = short_name_raw if (short_name_raw.startswith("«") or short_name_raw.startswith('"')) else f"«{short_name_raw}»"

    # Дополнительные условия (до 11 строк для юрлица)
    add_cond_raw = data.get("additional_conditions", "") or ""
    add_cond_lines = [l.strip() for l in add_cond_raw.splitlines() if l.strip()]
    add_cond = {f"ADD_COND_{i+1}": (f"- {add_cond_lines[i]}" if i < len(add_cond_lines) else "-") for i in range(11)}

    # INN/KPP форматы
    inn = c.get("inn", "")
    kpp = c.get("kpp", "")
    inn_kpp = f"{inn} / {kpp}" if kpp else inn
    inn_kpp_display = f"ИНН/КПП {inn} / {kpp}" if kpp else f"ИНН {inn}"

    # ФИО директора: краткая форма "Костоусов М.А."
    director_full = c.get("director_name", "")
    director_short = _short_director(director_full)

    return {
        "CONTRACT_NUMBER":              number,
        "CONTRACT_DATE":                format_date_legal(date_str),
        "COMPANY_LEGAL_FORM":           legal_form,
        "COMPANY_SHORT_NAME_DISPLAY":   short_name_display,
        "DIRECTOR_TITLE_AND_NAME":      f"{c.get('director_title','Генерального директора')}  {director_full}",
        "BASIS_DOCUMENT":               c.get("basis_document", "действующего на основании Устава"),
        "COMPANY_FULL_NAME":            full_name,
        "LEGAL_ADDRESS":                c.get("legal_address", ""),
        "POSTAL_ADDRESS":               c.get("postal_address", ""),
        "COMPANY_PHONE":                c.get("phone", ""),
        "INN_KPP":                      inn_kpp,
        "INN_KPP_DISPLAY":              inn_kpp_display,
        "OGRN":                         c.get("ogrn", ""),
        "BANK_ACCOUNT":                 c.get("bank_account", ""),
        "BANK_NAME":                    c.get("bank_name", ""),
        "CORRESPONDENT_ACCOUNT":        c.get("correspondent_account", ""),
        "BIK":                          c.get("bik", ""),
        "DIRECTOR_TITLE_SHORT":         c.get("director_title_short", "Генеральный директор"),
        "DIRECTOR_SHORT_NAME":          director_short,
        "EMPLOYEE_1_NAME":              e1.get("name", ""),
        "EMPLOYEE_1_DOB":               e1.get("dob", ""),
        "EMPLOYEE_1_PASSPORT":          e1.get("passport", ""),
        "EMPLOYEE_1_ISSUE_DATE":        e1.get("issue_date", ""),
        "EMPLOYEE_1_VALID_UNTIL":       e1.get("valid_until", ""),
        "EMPLOYEE_2_NAME":              e2.get("name", ""),
        "EMPLOYEE_2_DOB":               e2.get("dob", ""),
        "EMPLOYEE_2_PASSPORT":          e2.get("passport", ""),
        "EMPLOYEE_2_ISSUE_DATE":        e2.get("issue_date", ""),
        "EMPLOYEE_2_VALID_UNTIL":       e2.get("valid_until", ""),
        "COUNTRY_CITY":                 f"{data.get('country','')} {data.get('city','')}".strip(", "),
        "HOTEL":                        data.get("hotel", ""),
        "CHECK_IN_DATE":                data.get("check_in_date", ""),
        "CHECK_OUT_DATE":               data.get("check_out_date", ""),
        "NIGHTS":                       str(data.get("nights", "")),
        "ROOM_TYPE":                    data.get("room_type", ""),
        "ROOM_COUNT":                   data.get("room_count", ""),
        "INSURANCE":                    "Да" if data.get("insurance") else "Нет",
        **add_cond,
        "TOTAL_PRICE_RUB":              f"{total:,.0f} руб.".replace(",", " "),
    }


def _short_director(full_name: str) -> str:
    """'Костоусов Максим Александрович' → 'Костоусов М.А.'"""
    parts = full_name.strip().split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    if len(parts) == 2:
        return f"{parts[0]} {parts[1][0]}."
    return full_name


# ── Генерация документов ──────────────────────────────────────────────────────

def generate_individual_contract(data: dict) -> tuple[bytes, bytes]:
    """Возвращает (pdf_bytes, docx_bytes)."""
    replacements = build_individual_replacements(data)
    template = TEMPLATES_DIR / "contract_individual.docx"
    docx_bytes = _fill_document(template, replacements)
    pdf_bytes  = _docx_to_pdf(docx_bytes)
    return pdf_bytes, docx_bytes


def generate_legal_contract(data: dict) -> tuple[bytes, bytes]:
    """Возвращает (pdf_bytes, docx_bytes)."""
    replacements = build_legal_replacements(data)
    template = TEMPLATES_DIR / "contract_legal.docx"
    docx_bytes = _fill_document(template, replacements)
    pdf_bytes  = _docx_to_pdf(docx_bytes)
    return pdf_bytes, docx_bytes


def format_individual_filename(contract_number: str, client_name: str) -> str:
    return f"Договор № {contract_number} {_short_name(client_name)}"


def format_legal_filename(contract_number: str, company_name: str) -> str:
    return f"Договор № {contract_number} {company_name[:30].strip()}"


def _short_name(full_name: str) -> str:
    parts = full_name.strip().split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    if len(parts) == 2:
        return f"{parts[0]} {parts[1][0]}."
    return full_name
